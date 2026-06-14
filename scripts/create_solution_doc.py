# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── עמוד A4 + שוליים ──────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Inches(8.27)
section.page_height   = Inches(11.69)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ── פלטת צבעים ────────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1A, 0x37, 0x6C)
ACCENT_BLUE = RGBColor(0x26, 0x74, 0xC1)
MID_GREY    = RGBColor(0x59, 0x59, 0x59)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
PLACEHOLDER = RGBColor(0xAA, 0xAA, 0xAA)
GOLD        = RGBColor(0x8B, 0x60, 0x00)
GREEN_DARK  = RGBColor(0x1B, 0x5E, 0x20)
RED_DARK    = RGBColor(0x8B, 0x00, 0x00)
ORANGE      = RGBColor(0xE6, 0x5C, 0x00)

HEB_FONT = 'Arial'

# ══════════════════════════════════════════════════════════════════════════════
#  פונקציות עזר
# ══════════════════════════════════════════════════════════════════════════════
def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def shade_paragraph(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_rtl(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    tcPr.append(bidi)
    for para in cell.paragraphs:
        set_rtl(para)

def add_run(para, text, bold=False, italic=False, size=11, color=None, ltr=False, rtl=False):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.size = Pt(size)
    run.font.name = HEB_FONT
    rPr = run._r.get_or_add_rPr()
    # w:rtl חייב להופיע לפני w:lang לפי סדר ה-schema של OOXML
    if rtl:
        rtl_el = OxmlElement('w:rtl')
        rtl_el.set(qn('w:val'), '1')
        rPr.append(rtl_el)
    lang = OxmlElement('w:lang')
    lang.set(qn('w:bidi'), 'he-IL')
    rPr.append(lang)
    if color:
        run.font.color.rgb = color
    return run

def p(doc, text='', bold=False, italic=False, size=11, color=None,
      sb=0, sa=5, shade=None, rtl=True):
    para = doc.add_paragraph()
    if rtl:
        set_rtl(para)
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    if shade:
        shade_paragraph(para, shade)
    if text:
        add_run(para, text, bold=bold, italic=italic, size=size, color=color)
    return para

def hr(doc):
    para = doc.add_paragraph()
    set_rtl(para)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bot)
    pPr.append(pBdr)

# ── בלוקים מסוגננים ────────────────────────────────────────────────────────────
def part_banner(doc, text):
    para = p(doc, f'  {text}', bold=True, size=14, color=WHITE,
             sb=10, sa=2, shade='1A376C')

def q_header(doc, qn_str, title, hex_col='2674C1'):
    p(doc, f'  {qn_str}. {title}', bold=True, size=13, color=WHITE,
      sb=12, sa=4, shade=hex_col)

def sub(doc, text):
    p(doc, text, bold=True, size=11, color=DARK_BLUE, sb=10, sa=1)
    hr(doc)

def answer(doc, text, size=11):
    """שורת תשובה — טקסט מלא."""
    para = p(doc, text, size=size, color=MID_GREY, sb=3, sa=3)

def answer_bullet(doc, text, size=11, color=None):
    """נקודת תשובה עם •"""
    para = doc.add_paragraph()
    set_rtl(para)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after  = Pt(6)
    para.paragraph_format.left_indent  = Cm(0.5)
    add_run(para, '• ', bold=True, size=size, color=ACCENT_BLUE)
    add_run(para, text, size=size, color=color or MID_GREY)

def data_box(doc, lines):
    """תיבת ממצאים (צהובה) עם שורות נתונים."""
    p(doc, '  ממצאים מהנתונים:', bold=True, size=10, color=GOLD,
      sb=6, sa=0, shade='FFF8E1')
    for line in lines:
        para = doc.add_paragraph()
        set_rtl(para)
        shade_paragraph(para, 'FFFDE7')
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(3)
        para.paragraph_format.left_indent  = Cm(0.4)
        add_run(para, line, size=10, color=RGBColor(0x5D, 0x40, 0x00))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def insight(doc, text):
    """תיבת תובנה (ירוקה)."""
    p(doc, '  תובנה עסקית:', bold=True, size=10, color=GREEN_DARK,
      sb=6, sa=0, shade='E8F5E9')
    para = doc.add_paragraph()
    set_rtl(para)
    shade_paragraph(para, 'F1F8F1')
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(6)
    para.paragraph_format.left_indent  = Cm(0.4)
    add_run(para, text, size=10, color=GREEN_DARK)

def warning_box(doc, text):
    """תיבת אזהרה/ממצא חריג (אדום)."""
    p(doc, '  ממצא חריג — דורש בדיקה:', bold=True, size=10, color=RED_DARK,
      sb=6, sa=0, shade='FFEBEE')
    para = doc.add_paragraph()
    set_rtl(para)
    shade_paragraph(para, 'FFF3F3')
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(6)
    para.paragraph_format.left_indent  = Cm(0.4)
    add_run(para, text, size=10, color=RED_DARK)

def code_block(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(6)
    shade_paragraph(para, '1E1E1E')
    run = para.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)

CHARTS_DIR = r'C:\Users\Yosef\PycharmProjects\HAAT_assignment\charts'

def add_chart(doc, filename, width_cm=14.5, caption=None):
    import os
    para = doc.add_paragraph()
    set_rtl(para)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(4)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(os.path.join(CHARTS_DIR, filename), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        set_rtl(cp)
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(10)
        add_run(cp, caption, italic=True, size=9, color=MID_GREY, rtl=True)

def placeholder_section(doc):
    p(doc, '  [ תוכן השאלה יועתק לכאן לאחר קבלת שאר המטלה ]',
      italic=True, size=10, color=PLACEHOLDER, sb=4, sa=4, shade='F5F5F5')
    for _ in range(6):
        para = doc.add_paragraph()
        set_rtl(para)
        shade_paragraph(para, 'F0F4FF')
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)

# ══════════════════════════════════════════════════════════════════════════════
#  עמוד שער
# ══════════════════════════════════════════════════════════════════════════════
cv = doc.add_paragraph()
set_rtl(cv)
cv.alignment = WD_ALIGN_PARAGRAPH.CENTER
cv.paragraph_format.space_before = Pt(60)
cv.paragraph_format.space_after  = Pt(0)
shade_paragraph(cv, '1A376C')
add_run(cv, 'HAAT Delivery', bold=True, size=30, color=WHITE)

cv2 = doc.add_paragraph()
set_rtl(cv2)
cv2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cv2.paragraph_format.space_after = Pt(0)
shade_paragraph(cv2, '1A376C')
add_run(cv2, 'מטלת בית — אנליסט נתונים', size=18, color=RGBColor(0xAD, 0xD8, 0xE6))

cv3 = doc.add_paragraph()
set_rtl(cv3)
cv3.alignment = WD_ALIGN_PARAGRAPH.CENTER
cv3.paragraph_format.space_after = Pt(30)
shade_paragraph(cv3, '1A376C')
add_run(cv3, 'מחלקת אנליטיקה · 2026', size=13, color=RGBColor(0xCC, 0xCC, 0xCC))

info_t = doc.add_table(rows=3, cols=2)
info_t.style = 'Table Grid'
for i, (lbl, val) in enumerate([
        ('שם המגיש:',    '[שמך כאן]'),
        ('תאריך הגשה:', '[תאריך]'),
        ('כלים בשימוש:', 'Python · pandas · Jupyter Notebook')]):
    lc, rc = info_t.rows[i].cells
    shade_cell(lc, 'E8EDF5'); set_cell_rtl(lc)
    shade_cell(rc, 'FFFFFF'); set_cell_rtl(rc)
    add_run(lc.paragraphs[0], lbl, bold=True, size=11, color=DARK_BLUE)
    add_run(rc.paragraphs[0], val, size=11, color=MID_GREY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  תוכן עניינים
# ══════════════════════════════════════════════════════════════════════════════
p(doc, '  תוכן עניינים', bold=True, size=16, color=WHITE, shade='2674C1')
toc = [
    ('חלק א — הכרת העסק',               True),
    ('    שאלה 1: מבט ראשון על הנתונים', False),
    ('    שאלה 2: הגדרת מדדי הצלחה (KPI)', False),
    ('חלק ב — עומק תפעולי',               True),
    ('    שאלה 3: אמינות המשלוחים',        False),
    ('    שאלה 4: יעילות שליחים',          False),
    ('    שאלה 5: מהיכן מגיע הכסף',       False),
    ('חלק ג — חשיבה עסקית',               True),
    ('    שאלה 6: שימור לקוחות',           False),
    ('    שאלות 7–8: ישלים בהמשך',         False),
    ('שאלה 10: שאלת בונוס',               True),
]
for text, bold in toc:
    p(doc, text, bold=bold, size=11,
      color=DARK_BLUE if bold else MID_GREY, sb=0, sa=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  חלק א — הכרת העסק
# ══════════════════════════════════════════════════════════════════════════════
part_banner(doc, 'חלק א — הכרת העסק')

# ════════════════════════════════════════════════════════════════
#  שאלה 1 — מבט ראשון על הנתונים
# ════════════════════════════════════════════════════════════════
q_header(doc, 'שאלה 1', 'מבט ראשון על הנתונים')
p(doc, 'הנחיה: סקור את מסד הנתונים וכתוב סיכום קצר של מה שאתה רואה.',
  italic=True, size=10, color=MID_GREY, sb=0, sa=6)

# ── 1א: כמות ופיזור ────────────────────────────────────────────
sub(doc, 'א | כמות הזמנות — פיזור לפי חודשים ואזורים')

code_block(doc,
'df["Month"] = df["OrderDate"].dt.to_period("M")\n'
'df.groupby("Month").size()\n'
'df.groupby("AreaName").size().sort_values(ascending=False)')

data_box(doc, [
    'סה"כ הזמנות: 28,645  |  תקופה: פברואר–מאי 2024 (4 חודשים)',
    'פיזור חודשי: פבר\' 6,973 (24.3%)  |  מרץ 6,771 (23.6%)  |  אפר\' 7,455 (26.0%)  |  מאי 7,446 (26.0%)',
    'פיזור אזורי: Umm al-Fahem 8,784 (30.7%)  |  Sakhnin 4,569 (16.0%)  |  Kfar Qaree-Arara 3,709 (12.9%)',
    '4 האזורים הגדולים = 70.9% מהנפח הכולל  |  3 האזורים הקטנים = 4.2% בלבד',
    'האזורים הקטנים: Kafr Qasem 461  |  Jisr az-Zarqa 420  |  Ramla 321',
])

answer(doc, 'מסד הנתונים מכיל 28,645 הזמנות שבוצעו ב-11 אזורים שונים על פני 4 חודשים (פברואר–מאי 2024).')
answer_bullet(doc, 'נפח חודשי: יציב יחסית, עם קפיצה קלה של כ-10% ממרץ לאפריל ואז התייצבות. אין עלייה דרמטית שמעידה על צמיחה חזקה.')
answer_bullet(doc, 'פיזור אזורי: מרוכז מאוד — 4 ערים גדולות (Umm al-Fahem, Sakhnin, Kfar Qaree-Arara, Baqa al-Gharbiyye) מהוות ~71% מהפעילות. שאר 7 האזורים מחלקים ביניהם 29% בלבד.')
answer_bullet(doc, 'Umm al-Fahem בלבד אחראית על כמעט שליש (30.7%) מכלל ההזמנות — ברור שהיא השוק הראשי של הפלטפורמה.')
insight(doc, 'הריכוזיות הגבוהה ב-4 ערים יוצרת תלות — כל שיבוש תפעולי בUmm al-Fahem ישפיע מיידית על ~31% מהמכירות. מנגד, יש פוטנציאל צמיחה ברור ב-7 האזורים הקטנים שעדיין מתחת ל-10% מהנפח.')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── 1ב: ערך הזמנה ותשלום ───────────────────────────────────────
sub(doc, 'ב | ממוצע ערך הזמנה ושיטת תשלום נפוצה')

code_block(doc,
'df["Price"].describe()\n'
'df["PaymentMethod"].value_counts()  # 0=מזומן, 1=אשראי')

data_box(doc, [
    'ממוצע ערך הזמנה: 118.96 ₪  |  חציון: 102.80 ₪  |  סטיית תקן: 79.57 ₪',
    'טווח מחירים: 1.30 ₪ (מינימום) עד 2,743.60 ₪ (מקסימום)',
    '75% מהזמנות: עד 150.70 ₪  |  5 הזמנות בלבד מעל 1,000 ₪ (כנראה הזמנות קבוצתיות)',
    'שיטת תשלום: מזומן — 19,626 (68.5%)  |  אשראי — 9,019 (31.5%)',
    'אין ערכים >1 בעמודת PaymentMethod — רק מזומן ואשראי בפועל',
])

answer(doc, 'ממוצע ערך ההזמנה הוא 118.96 ₪, עם חציון נמוך יותר של 102.80 ₪ — הפרש זה מעיד על כך שמיעוט של הזמנות גדולות מושך את הממוצע כלפי מעלה (התפלגות מוטה ימינה).')
answer_bullet(doc, 'שיטת התשלום הנפוצה ביותר היא מזומן — 68.5% מהעסקאות. זאת תכונה ייחודית לשוק זה, שכן בפלטפורמות משלוח אחרות אשראי/דיגיטל הוא הנפוץ.')
answer_bullet(doc, 'הפרש בין ממוצע לחציון (16 ₪) קטן יחסית, אך סטיית תקן של 79.57 ₪ מראה שיש פיזור רחב — חלק מהלקוחות מזמינים הרבה, רוב ההזמנות הן בסביבות 60–150 ₪.')
insight(doc, 'שיעור מזומן גבוה של 68.5% מקשה על ניתוח התנהגות לקוח ועל מניעת הונאות. כדאי לבדוק האם ניתן לעודד לקוחות לעבור לתשלום דיגיטלי — זה ישפר מעקב ויאפשר תכניות נאמנות.')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── 1ג: זמן שיוך נהג ───────────────────────────────────────────
sub(doc, 'ג | זמן ממוצע עד שיוך נהג (OrderDate → DriverCandidateAssignedDate)')

code_block(doc,
'df["time_to_assign_min"] = (\n'
'    (df["DriverCandidateAssignedDate"] - df["OrderDate"])\n'
'    .dt.total_seconds() / 60\n'
')\n'
'df["time_to_assign_min"].describe()')

data_box(doc, [
    'הזמנות שקיבלו נהג: 18,824 מתוך 28,645 (65.7%)  |  34.3% לא קיבלו נהג כלל',
    'ממוצע: 8.4 דקות  |  חציון: 6.0 דקות  |  P25: 2.4 דקות  |  P75: 11.8 דקות',
    '10% מהלקוחות שקיבלו נהג — המתינו יותר מ-19.2 דקות (P90)',
    'מקסימום: 58.6 דקות — ממתין כמעט שעה לשיוך נהג',
])

answer(doc, 'מבין 28,645 ההזמנות, רק 18,824 (65.7%) קיבלו נהג משויך. מבין אלה, זמן השיוך החציוני הוא 6 דקות — נתון סביר. אולם:')
answer_bullet(doc, 'חציון: 6.0 דקות — חצי מהלקוחות שמקבלים נהג ממתינים עד 6 דקות.')
answer_bullet(doc, '10% ממתינים מעל 19 דקות (P90) — זה כבר גבולי לחוויית לקוח טובה.')
answer_bullet(doc, 'קריטי: 34.3% מהזמנות לעולם לא קיבלו נהג — הזמן הממוצע שלהן לשיוך הוא אינסוף. אם נכלול אותן בחישוב, ממוצע ה"זמן לשיוך" בפועל הוא גבוה בהרבה.')
insight(doc, 'ה-6 דקות החציוניות נשמעות טוב, אבל הן מייצגות רק 65.7% מהלקוחות. השאלה האמיתית היא: מדוע 34.3% מהזמנות לא מצאו נהג? זה הממצא הקריטי ביותר בשאלה זו.')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── 1ד: ממצא חריג ──────────────────────────────────────────────
sub(doc, 'ד | ממצא חריג — דבר אחד שנראה חשוד ומצריך בדיקה')

data_box(doc, [
    'עמודה: ArriveDate (ואיתה DriverCandidateAssignedDate, PickedDate)',
    '9,973 הזמנות (34.8%) ללא ArriveDate = לא נמסרו ללקוח לעולם',
    '9,679 מתוכן (97.1%) — ללא נהג משויך בכלל (DriverCandidateAssignedDate = NULL)',
    '294 הזמנות — קיבלו נהג אך בכל זאת לא נמסרו',
    '1,109 הזמנות נדחו ע"י המסעדה (3.9%) — חלק מהסיבות לאי-מסירה',
])

warning_box(doc,
    'כמעט 1 מכל 3 הזמנות (34.8%) לא הגיעה ללקוח. '
    'זוהי תוצאה עסקית כשלונית חמורה: לקוח שהזמין, ציפה — ולא קיבל את ההזמנה. '
    'ב-97.1% מהמקרים הסיבה היא היעדר נהג זמין. '
    'הפלטפורמה לוקה בחסר צד-היצע (סאפליי סייד) קריטי.')

answer(doc, 'הממצא החריג ביותר הוא שיעור אי-המסירה: 34.8% מההזמנות (9,973 הזמנות) לא נמסרו ללקוח.')
answer_bullet(doc, 'עמודת ArriveDate: NULL ב-9,973 רשומות — מעיד שהנהג מעולם לא הגיע ללקוח.')
answer_bullet(doc, 'הסיבה העיקרית: 9,679 מהן לא קיבלו נהג כלל (DriverCandidateAssignedDate = NULL), כלומר הבעיה היא מחסור בנהגים — לא כשל בשלב המשלוח עצמו.')
answer_bullet(doc, 'זמן ומיקום: יש לבדוק האם הבעיה מרוכזת בשעות שיא (ערב) או באזורים ספציפיים — זה יכוון לפתרון.')
insight(doc, 'שיעור אי-מסירה של 34.8% הוא גבוה לכל סטנדרט בתחום. פלטפורמות בוגרות כגון וולט/אובר איטס עובדות עם ~5-10% אי-מסירה. הפער מצביע על בעיית ניהול צי נהגים (פליט מנג\'מנט) חמורה שצריכה לעמוד בראש סדר העדיפויות.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  שאלה 2 — מדדי הצלחה (KPI Framework)
# ════════════════════════════════════════════════════════════════
q_header(doc, 'שאלה 2', 'הגדרת מדדי הצלחה (KPI Framework)')
p(doc,
  'הנחיה: הגדר את 5 המדדים החשובים ביותר לפלטפורמת משלוחים. '
  'לכל מדד: נוסחה מדויקת, מה הוא מודד ולמה, ערך תקין/מדאיג, צוות אחראי. '
  'חובה: שיעור הצלחת משלוח · מהירות שיוך נהג · מדד מסחרי.',
  italic=True, size=10, color=MID_GREY, sb=0, sa=8)

# ─── טבלת KPI ───────────────────────────────────────────────────
kpi_data = [
    {
        'num': '1',
        'name': 'שיעור הצלחת משלוח\n(דליוורי סאקסס רייט)',
        'formula': 'COUNT(ArriveDate IS NOT NULL)\n÷ COUNT(*) × 100',
        'what': 'אחוז ההזמנות שהגיעו ללקוח בפועל. '
                'זהו המדד הבסיסי ביותר — אם המשלוח לא הגיע, '
                'כל שאר המדדים אינם רלוונטיים.',
        'good': 'תקין: >90%\nמדאיג: <80%\nכרגע: 65.2% — חריג!',
        'team': 'תפעול'
    },
    {
        'num': '2',
        'name': 'מהירות שיוך נהג\n(קוריאר אסיינמנט טיים)',
        'formula': 'MEDIAN(\n  DriverCandidateAssignedDate\n  − OrderDate\n) בדקות\n[רק הזמנות עם נהג]',
        'what': 'כמה מהר נהג מוקצה להזמנה מרגע שהונחה. '
                'מדד זה מייצג את יעילות מוקד השיגור ואת זמינות הנהגים.',
        'good': 'תקין: חציון <5 דק\'\nמדאיג: >15 דק\'\nכרגע: 6 דק\' — גבולי',
        'team': 'תפעול'
    },
    {
        'num': '3',
        'name': 'ממוצע ערך הזמנה\n(אוורג\' אורדר וואלו)',
        'formula': 'SUM(Price)\n÷ COUNT(*)',
        'what': 'הכנסה ממוצעת לכל הזמנה. '
                'מדד מסחרי מרכזי — ירידה עשויה להעיד על שחיקת מחירים '
                'או על שינוי בהרכב המוצרים.',
        'good': 'תקין: >120 ₪ ועולה\nמדאיג: <80 ₪ או ירידה\nכרגע: 118.96 ₪',
        'team': 'פיננסים / מסחר'
    },
    {
        'num': '4',
        'name': 'שיעור דחייה ע"י מסעדה\n(רסטורנט ריג\'קשן רייט)',
        'formula': 'COUNT(RejectedDateByRestaurant\n  IS NOT NULL)\n÷ COUNT(*) × 100',
        'what': 'אחוז ההזמנות שהמסעדה סירבה לקבל. '
                'מדד אמינות-צד-ספק — שיעור גבוה פוגע בחוויית הלקוח '
                'ומצביע על שותפויות בעייתיות.',
        'good': 'תקין: <2%\nמדאיג: >5%\nכרגע: 3.9%',
        'team': 'תפעול / שותפויות'
    },
    {
        'num': '5',
        'name': 'שיעור כיסוי נהגים\n(דרייבר קוברג\' רייט)',
        'formula': 'COUNT(DriverCandidateAssignedDate\n  IS NOT NULL)\n÷ COUNT(*) × 100',
        'what': 'אחוז ההזמנות שקיבלו נהג משויך. '
                'מדד זמינות הצי — שיעור נמוך מעיד על מחסור בנהגים '
                'ומסביר ישירות את שיעור אי-המסירה הגבוה.',
        'good': 'תקין: >90%\nמדאיג: <80%\nכרגע: 65.7% — חריג!',
        'team': 'תפעול / ניהול צי'
    },
]

kpi_table = doc.add_table(rows=1, cols=6)
kpi_table.style = 'Table Grid'
col_w    = [Cm(0.9), Cm(3.0), Cm(2.9), Cm(4.0), Cm(2.4), Cm(2.2)]
col_hdrs = ['#', 'שם המדד', 'נוסחה', 'מה זה מודד ולמה', 'ערך תקין / מדאיג\nמצב נוכחי', 'צוות\nאחראי']

# header row
hrow = kpi_table.rows[0]
for i, (hdr, w) in enumerate(zip(col_hdrs, col_w)):
    cell = hrow.cells[i]
    cell.width = w
    shade_cell(cell, '1A376C')
    set_cell_rtl(cell)
    add_run(cell.paragraphs[0], hdr, bold=True, size=9, color=WHITE)

# data rows
shade_alt = ['FFFFFF', 'F0F4FF']
for kpi in kpi_data:
    row   = kpi_table.add_row()
    shade = shade_alt[int(kpi['num']) % 2]
    vals  = [kpi['num'], kpi['name'], kpi['formula'], kpi['what'], kpi['good'], kpi['team']]
    for i, (val, w) in enumerate(zip(vals, col_w)):
        cell = row.cells[i]
        cell.width = w
        set_cell_rtl(cell)
        is_num = (i == 0)
        shade_cell(cell, 'E8EDF5' if is_num else shade)
        clr = DARK_BLUE if is_num else (
              RED_DARK if 'חריג' in val else MID_GREY)
        add_run(cell.paragraphs[0], val,
                bold=(i <= 1), size=9, color=clr)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# נימוק הבחירה
sub(doc, 'נימוק בחירת 5 המדדים')

answer(doc,
    'הבחירה מבוססת על לוגיקה של "משפך הערך" — כל מדד מייצג שלב קריטי במסע ההזמנה:')
answer_bullet(doc,
    'מדד 1 (שיעור הצלחת משלוח) — תוצאה סופית. '
    'אם זה נכשל, כל השאר לא משנה. חובת המטלה.',
    color=DARK_BLUE)
answer_bullet(doc,
    'מדד 2 (מהירות שיוך נהג) — תהליך. '
    'מסביר חלק מהכשלונות ומודד יעילות שיגור. חובת המטלה.',
    color=DARK_BLUE)
answer_bullet(doc,
    'מדד 3 (AOV) — כלכלה. '
    'ללא הכנסה בריאה לכל הזמנה, העסק אינו ברי-קיימא. חובת המטלה.',
    color=DARK_BLUE)
answer_bullet(doc,
    'מדד 4 (דחייה ע"י מסעדה) — אמינות-ספק. '
    '3.9% מהזמנות נדחות — כל דחייה פוגעת בלקוח ובמוניטין. '
    'חיוני לאיכות השירות.',
    color=DARK_BLUE)
answer_bullet(doc,
    'מדד 5 (כיסוי נהגים) — יכולת-אספקה. '
    '34.3% מהזמנות ללא נהג — זאת הבעיה הכי חמורה שמתגלה מהנתונים, '
    'וחייבים לעקוב אחריה.',
    color=DARK_BLUE)

insight(doc,
    'שלושת המדדים שמראים אות אזהרה עכשיו: '
    'שיעור הצלחה 65.2%, כיסוי נהגים 65.7%, ושיעור דחיית מסעדות 3.9%. '
    'מדד ממוצע ערך ההזמנה (118.96 ₪) קרוב לגבול התחתון של "תקין" — '
    'יש לעקוב האם הוא יורד או עולה לאורך זמן.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  חלק ב — עומק תפעולי
# ══════════════════════════════════════════════════════════════════════════════
part_banner(doc, 'חלק ב — עומק תפעולי')

# ── שאלה 3 ──────────────────────────────────────────────────────
q_header(doc, 'שאלה 3', 'אמינות המשלוחים')
p(doc, 'הנחיה: נתח ביצועי משלוח תוך שימוש ב-OrderDate, ArriveDate, DriverCandidateAssignedDate ו-RejectedDateByRestaurant.',
  italic=True, size=10, color=MID_GREY, sb=0, sa=6)

sub(doc, 'א | שיעור הזמנות שנמסרו — לפי אזור ולפי חודש')

code_block(doc,
'df["delivered"] = df["ArriveDate"].notna()\n'
'area_rate    = df.groupby("AreaName")["delivered"].agg(["count","sum"]).assign(rate=lambda x: x["sum"]/x["count"]*100)\n'
'monthly_rate = df.groupby("Month")["delivered"].mean() * 100')

data_box(doc, [
    'שיעור מסירה כולל: 65.2%  |  סה"כ הזמנות: 28,645  |  נמסרו: 18,672',
    '',
    'לפי חודש:',
    '  פברואר 2024: 6,973 הזמנות | 4,531 נמסרו | 65.0%',
    '  מרץ 2024:    6,771 הזמנות | 4,728 נמסרו | 69.8%  ← שיא',
    '  אפריל 2024:  7,455 הזמנות | 4,649 נמסרו | 62.4%  ← נפילה',
    '  מאי 2024:    7,446 הזמנות | 4,764 נמסרו | 64.0%',
    '',
    'לפי אזור (מיון יורד):',
    '  Kafr Qasem:        461 הזמנות | 97.0%  ▲▲',
    '  Jisr az-Zarqa:     420 הזמנות | 96.2%  ▲▲',
    '  Lod:               637 הזמנות | 94.8%  ▲▲',
    '  Ramla:             321 הזמנות | 91.6%  ▲▲',
    '  Sakhnin:         4,569 הזמנות | 75.7%  ▲',
    '  Nazareth:        2,948 הזמנות | 74.5%  ▲',
    '  Tamra - Kabul:   2,104 הזמנות | 65.4%  ~',
    '  Umm al-Fahem:    8,784 הזמנות | 58.6%  ▼',
    '  Shefa-Amr:       1,493 הזמנות | 58.3%  ▼',
    '  Baqa al-Gharbiyye: 3,199 הזמנות | 57.0%  ▼',
    '  Kfar Qaree - Arara: 3,709 הזמנות | 55.4%  ▼▼',
])

answer(doc, 'שיעור המסירה הכולל עומד על 65.2% — נמוך מאוד לפי כל סטנדרט בתחום. הנתונים החודשיים מגלים שאין שיפור מגמתי: מרץ הטוב ביותר (69.8%) ואפריל ירד חזרה ל-62.4%, ומאי מתייצב על 64%.')
answer_bullet(doc, 'אין מגמת שיפור ברורה לאורך 4 החודשים — הבעיה מבנית ולא עונתית.')
answer_bullet(doc, 'הפיזור האזורי מדהים: 4 הערים הקטנות (Kafr Qasem, Jisr az-Zarqa, Lod, Ramla) בין 91% ל-97%, בעוד 4 הערים הגדולות בין 55% ל-58%.')
insight(doc, 'הפער בין ערים קטנות (97%) לגדולות (55-58%) הוא מוחלט. לא מדובר בהבדל קטן — מדובר בפלטפורמה שונה לחלוטין מבחינת חוויית לקוח. כנראה שבערים הקטנות נפח ההזמנות נמוך מספיק כדי שהנהגים הזמינים יכסו אותן, ואילו בערים הגדולות — בייחוד Umm al-Fahem — ביקוש הנהגים עולה בהרבה על ההיצע.')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ב | נפח הזמנות ושיעור משלוח לפי שעה — זיהוי תבניות')

code_block(doc,
'df["Hour"] = df["OrderDate"].dt.hour\n'
'hourly = df.groupby("Hour").agg(\n'
'    orders        = ("Id", "count"),\n'
'    delivery_rate = ("delivered", "mean")\n'
').assign(delivery_rate = lambda x: x["delivery_rate"] * 100)')

data_box(doc, [
    'שעות שיא נפח: 20:00 (2,991) | 21:00 (2,953) | 22:00 (2,593)',
    'שיעור מסירה הטוב ביותר: 06:00 (97.1%) — אך רק 34 הזמנות, לא מייצג',
    'שיעור מסירה הגרוע ביותר: 18:00 (60.2%)',
    '',
    'שעות ערב — נפח גבוה + מסירה נמוכה:',
    '  17:00 | 2,004 הזמנות | 65.5%',
    '  18:00 | 2,110 הזמנות | 60.2%  ← שפל',
    '  19:00 | 2,363 הזמנות | 65.7%',
    '  20:00 | 2,991 הזמנות | 64.2%',
    '  21:00 | 2,953 הזמנות | 61.0%',
    '  22:00 | 2,593 הזמנות | 62.6%',
    '',
    'שעות לילה/בוקר — נפח נמוך + מסירה גבוהה יותר:',
    '  02:00 | 488 הזמנות | 70.9%',
    '  03:00 | 272 הזמנות | 73.2%',
    '  04:00 | 107 הזמנות | 74.8%',
])

answer(doc, 'ישנה תבנית ברורה: ככל שנפח ההזמנות עולה לאורך היום — שיעור המסירה יורד. הוכחה חזקה לבעיית מחסור בנהגים בשעות שיא.')
answer_bullet(doc, 'שעות 17:00–22:00 (שעות הערב): שיאי הנפח עם שיעורי מסירה של 60-65% בלבד. בשעות אלה הביקוש לנהגים גדול מהיצע הנהגים הזמינים.')
answer_bullet(doc, 'שעות 02:00–08:00 (לילה/בוקר): נפח נמוך (28-488 הזמנות) עם שיעורי מסירה של 71-97%. כשיש מעט הזמנות — הנהגים הזמינים מספיקים לכולן.')
answer_bullet(doc, 'שעה 18:00 היא הנקודה הגרועה ביותר: 60.2% — כנראה השעה שבה הביקוש קופץ (ארוחת ערב) אך צי הנהגים עדיין לא מלא.')
insight(doc, 'הממצא הזה הוא אחד החשובים בכל הניתוח: הבעיה אינה אחידה לאורך היום — היא מרוכזת בשעות ערב ספציפיות. פתרון ממוקד (גיוס נהגים לוורקשיפטים של 17:00–22:00) יכול לשפר דרמטית את שיעור המסירה הכולל.')
add_chart(doc, 'q3b_hourly.png', caption='תרשים 2 — נפח הזמנות (עמודות) ושיעור מסירה (קו) לפי שעה')
add_chart(doc, 'q3a_monthly.png', caption='תרשים 3 — שיעור מסירה לפי חודש')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ג | השוואת ביצועים בין אזורים')

data_box(doc, [
    'ממוצע כלל-פלטפורמה: 65.2%',
    '',
    'קבוצה א — ביצוע גבוה מאוד (ערים קטנות, נפח נמוך):',
    '  Kafr Qasem:     461 הזמנות | 97.0% | פער: +31.8%',
    '  Jisr az-Zarqa:  420 הזמנות | 96.2% | פער: +31.0%',
    '  Lod:            637 הזמנות | 94.8% | פער: +29.6%',
    '  Ramla:          321 הזמנות | 91.6% | פער: +26.4%',
    '',
    'קבוצה ב — ביצוע בינוני (ערים בינוניות):',
    '  Sakhnin:        4,569 הזמנות | 75.7% | פער: +10.5%',
    '  Nazareth:       2,948 הזמנות | 74.5% | פער: +9.3%',
    '  Tamra - Kabul:  2,104 הזמנות | 65.4% | פער: +0.2%',
    '',
    'קבוצה ג — ביצוע נמוך (ערים גדולות, נפח גבוה):',
    '  Umm al-Fahem:       8,784 הזמנות | 58.6% | פער: -6.6%',
    '  Shefa-Amr:          1,493 הזמנות | 58.3% | פער: -6.9%',
    '  Baqa al-Gharbiyye:  3,199 הזמנות | 57.0% | פער: -8.2%',
    '  Kfar Qaree - Arara: 3,709 הזמנות | 55.4% | פער: -9.8%',
])

answer(doc, 'האזורים מתחלקים לשלוש קבוצות ברורות על פי שיעור המסירה, ויש קשר הפוך ישיר בין נפח ההזמנות לבין איכות המסירה.')
answer_bullet(doc, 'ערים קטנות (321–637 הזמנות) = 91–97% מסירה. כמות ההזמנות נמוכה מספיק כדי שהנהגים הזמינים יכסו כמעט הכל.')
answer_bullet(doc, 'ערים גדולות (1,493–8,784 הזמנות) = 55–58% מסירה. ה-bottleneck הוא כיסוי נהגים — הביקוש גדל מהר יותר מהיצע הנהגים.')
answer_bullet(doc, 'Umm al-Fahem בולטת לרעה: הגדולה ביותר (30.7% מהנפח) אך שיעור מסירה של 58.6% בלבד — כישלון של כמעט כל הזמנה שנייה.')
answer_bullet(doc, 'Shefa-Amr חריגה: נפח בינוני (1,493) אך ביצועים דומים לערים הגדולות — ייתכן מחסור ספציפי בנהגים באזור זה.')
insight(doc, 'הממצא המרכזי: הפלטפורמה "עובדת" בערים הקטנות ו"נכשלת" בגדולות. פתרון האספקה (נהגים) חייב להיות מותאם לכל אזור בנפרד, ולא פתרון אחיד. עדיפות ראשונה: Kfar Qaree - Arara ו-Baqa al-Gharbiyye שמציגות את הביצועים הגרועים ביותר ביחס לנפחן.')
add_chart(doc, 'q3a_area.png', caption='תרשים 1 — שיעור מסירה לפי אזור (ירוק = טוב, אדום = בעייתי)')

doc.add_page_break()

# ── שאלה 4 ────────────────────────────────────────────────────────
q_header(doc, 'שאלה 4', 'יעילות שליחים (קוריאר אפישנסי)')
p(doc, 'הנחיה: השתמש בטבלת ההזמנות כדי לנתח ביצועי שליחים ברמת הפרט.',
  italic=True, size=10, color=MID_GREY, sb=0, sa=6)

sub(doc, 'א | דירוג שליחים לפי הזמנות שנמסרו — TOP 10 ו-BOTTOM 10')

code_block(doc,
'driver_stats = df[df["DriverId"].notna()].groupby("DriverId").agg(\n'
'    total_assigned  = ("Id", "count"),\n'
'    total_delivered = ("delivered", "sum"),\n'
'    first_order     = ("OrderDate", "min"),\n'
'    last_order      = ("OrderDate", "max"),\n'
').reset_index()\n'
'driver_stats.sort_values("total_delivered", ascending=False)')

data_box(doc, [
    'סה"כ שליחים פעילים (עם לפחות הזמנה אחת): 918',
    '',
    'TOP 10 — לפי כמות משלוחים שנמסרו:',
    '  שליח 918: 221 נמסרו | 221 שובצו | 100% | 12 ימים פעיל',
    '  שליח 472: 221 נמסרו | 221 שובצו | 100% | 12 ימים פעיל',
    '  שליח 738: 221 נמסרו | 221 שובצו | 100% | 12 ימים פעיל',
    '  שליח 538: 163 נמסרו | 163 שובצו | 100% | 121 ימים פעיל',
    '  שליח 161: 146 נמסרו | 146 שובצו | 100% | 118 ימים פעיל',
    '  שליח 875: 122 נמסרו | 123 שובצו | 99.2% | 121 ימים פעיל',
    '  שליח 365: 109 נמסרו | 109 שובצו | 100% | 119 ימים פעיל',
    '  שליח 334: 101 נמסרו | 101 שובצו | 100% | 120 ימים פעיל',
    '  שליח 188: 101 נמסרו | 102 שובצו | 99.0% | 121 ימים פעיל',
    '  שליח 265: 100 נמסרו | 100 שובצו | 100% | 119 ימים פעיל',
    '',
    'BOTTOM 10 — לפי כמות משלוחים שנמסרו:',
    '  שליחים 384, 391, 395, 409, 410, 425, 432, 434, 869, 5: 1 משלוח | 1 יום פעיל',
])

answer(doc, '3 השליחים הראשונים (918, 472, 738) בולטים מאוד: 221 משלוחים כל אחד ב-12 ימים בלבד עם שיעור הצלחה של 100%. המספר הזה חשוד — ראו סעיף ג.')
answer_bullet(doc, 'שליחים 4–10 בדירוג (538, 161, 875 וכולי) הם שליחים "אמיתיים": 100–163 משלוחים על פני 118–121 ימים — כלומר כ-1–1.3 משלוחים ביום בממוצע.')
answer_bullet(doc, 'ה-BOTTOM 10 הם שליחים חדשים ביומם הראשון — ערך נמוך מסיבה טבעית, לא כשל.')
add_chart(doc, 'q4a_ranking.png', caption='תרשים 4 — TOP 10 ו-BOTTOM 10 שליחים לפי כמות משלוחים שנמסרו')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ב | מדד יעילות הוגן — משלוחים מנורמלים ליום פעיל')

code_block(doc,
'driver_stats["active_days"] = (\n'
'    (driver_stats["last_order"] - driver_stats["first_order"]).dt.days + 1\n'
')\n'
'driver_stats["efficiency"] = (\n'
'    driver_stats["total_delivered"] / driver_stats["active_days"]\n'
')')

answer(doc, 'בעיית דירוג לפי כמות בלבד: שליח שפעיל 90 יום ייצבור יותר משלוחים מאחד שהתחיל אתמול — אך זה לא אומר שהוא יעיל יותר. נדרש מדד שמנרמל לזמן.')
p(doc, '  הגדרת מדד היעילות:', bold=True, size=10, color=DARK_BLUE, sb=6, sa=2)
answer(doc, 'יעילות שליח = סה"כ הזמנות שנמסרו ÷ מספר ימים פעילים')
answer(doc, 'ימים פעילים = (תאריך ההזמנה האחרונה − תאריך ההזמנה הראשונה) + 1')
answer_bullet(doc, 'מה המדד מודד: כמה משלוחים מוצלחים מבצע השליח בממוצע בכל יום שהוא פעיל — ללא קשר לתקופה הכוללת.')
answer_bullet(doc, 'למה הוגן: שליח חדש שפעיל יומיים ועשה 4 משלוחים (יעילות 2.0) מדורג גבוה יותר מאשר שליח ותיק שפעיל 100 ימים ועשה 50 משלוחים (יעילות 0.5).')
answer_bullet(doc, 'מגבלה: שליח שפעיל יום אחד בלבד עם 3 משלוחים מקבל 3.0 — גבוה מאוד אך לא בהכרח מייצג. לכן יש לשלב מינימום ימים פעילים בסינון.')
insight(doc, 'מדד מנורמל-ליום הוגן יותר מספירה גולמית, אך גם הוא לא מושלם. שליח שפועל רק בסופי שבוע יראה פחות ימים פעילים מאחד שפועל כל יום — גם אם שניהם עושים אותה כמות משלוחים לשעת עבודה. מדד אידיאלי ידרוש נתוני שעות עבודה.')
add_chart(doc, 'q4b_efficiency.png', caption='תרשים 5 — TOP 10 שליחים לפי מדד יעילות מנורמל (≥7 ימים ו-≥10 משלוחים)')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ג | התפלגות מדד היעילות ואיתור חריגים')

data_box(doc, [
    'סה"כ שליחים: 918',
    'ממוצע: 0.570 משלוחים/יום  |  חציון: 0.381 משלוחים/יום',
    'P75: 0.667  |  P90: 1.000  |  P95: 1.335  |  מקסימום: 18.417',
    '',
    'התפלגות לפי טווחים:',
    '  0.0 – 1.0 משלוחים/יום: 860 שליחים (93.7%) — הרוב המוחלט',
    '  1.0 – 2.0 משלוחים/יום:  51 שליחים (5.6%)',
    '  2.0 – 3.0 משלוחים/יום:   4 שליחים (0.4%)',
    '  10+       משלוחים/יום:   3 שליחים (0.3%) — חריגים קיצוניים!',
    '',
    'גבול חריגים (Q3 + 3×IQR): 1.987 משלוחים/יום',
    'מספר חריגים מעל הגבול: 35 שליחים',
    '',
    'החריגים הקיצוניים ביותר:',
    '  שליח 918: 18.42 משלוחים/יום | 221 נמסרו | 12 ימים פעיל',
    '  שליח 738: 18.42 משלוחים/יום | 221 נמסרו | 12 ימים פעיל',
    '  שליח 472: 18.42 משלוחים/יום | 221 נמסרו | 12 ימים פעיל',
])

warning_box(doc,
    '3 שליחים (918, 738, 472) ביצעו 221 משלוחים כל אחד ב-12 ימים בלבד עם 100% הצלחה — '
    'כלומר 18.4 משלוחים ביום. אם משמרת היא 8 שעות, מדובר במשלוח אחד כל 26 דקות ללא כל כשל. '
    'זה כמעט בלתי אפשרי פיזית לשליח יחיד.')

answer(doc, '35 שליחים חורגים מגבול האאוטליירים (1.987). מרביתם — שליחים חדשים עם 1–2 ימים פעילים ו-2–3 משלוחים, מה שמנפח את המדד באופן סטטיסטי.')
answer_bullet(doc, '3 השליחים הקיצוניים (918, 738, 472) הם ממצא אחר לגמרי: נפח עצום, תקופה קצרה, הצלחה מושלמת. הסבר סביר: חשבונות טסט שנוצרו לבדיקת המערכת, חשבון יחיד שמייצג צוות, או שגיאת נתונים בייבוא.')
answer_bullet(doc, 'ניתן לראות שהמספר המדויק 221 חוזר על עצמו פעמיים — ל-918, ל-738 וגם ל-472. הסיכוי שזה מקרי קטן מאוד, ומחזק את הסברת חשבונות הטסט.')
insight(doc, 'בניתוח יעילות שליחים אמיתי — יש לסנן מראש שליחים עם פחות מ-7 ימים פעילים ו-10 משלוחים, ולבדוק אם ה-3 חריגים הקיצוניים מופיעים בתשתית ה-CRM כחשבונות שליחים לגיטימיים.')
add_chart(doc, 'q4c_distribution.png', caption='תרשים 6 — התפלגות מדד יעילות השליחים (שמאל: כלל, ימין: זום 0–2)')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ד | נתונים נוספים הדרושים לניתוח יעילות מלא')

answer(doc, 'הנתונים הקיימים מאפשרים ניתוח ראשוני בלבד. לניתוח יעילות מלא נדרש:')
answer_bullet(doc, 'שעות עבודה בפועל: כדי לנרמל לשעות עבודה ולא לימים קלנדריים. שליח שעובד 4 שעות ביום לא ניתן להשוות לאחד שעובד 10 שעות.')
answer_bullet(doc, 'תאריך הצטרפות רשמי: אנחנו משתמשים בהזמנה הראשונה כתחליף, אך שליח שהצטרף ולא עבד שבוע ואז חזר — המדד שלנו ידלג על אותו שבוע ויראה אותו כמי שלא "פעיל" באותו זמן.')
answer_bullet(doc, 'אזור עבודה של כל שליח: שליח שעובד ב-Umm al-Fahem מתמודד עם נפח גדול בהרבה משליח ב-Ramla — ההשוואה ביניהם אינה הוגנת ללא קונטרול לאזור.')
answer_bullet(doc, 'דירוגי לקוחות לשליח: יעילות כמותית בלי איכות שירות היא מדד חלקי. שליח מהיר שמקבל דירוגים נמוכים גרוע משליח קצת איטי יותר שמקבל 5 כוכבים.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  שאלה 5 — מהיכן מגיע הכסף?
# ════════════════════════════════════════════════════════════════
q_header(doc, 'שאלה 5', 'מהיכן מגיע הכסף? (וור איז דה מאני קאמינג פרום)')
p(doc, 'הנחיה: נתח כיצד הכנסות HAAT מתפלגות לפי אזורים ועסקים.',
  italic=True, size=10, color=MID_GREY, sb=0, sa=6)

sub(doc, 'א | פירוק הכנסות לפי אזור ולפי עסק')

code_block(doc,
'total_rev = df["Price"].sum()\n'
'area_rev  = df.groupby("AreaName").agg(\n'
'    revenue=("Price","sum"), orders=("Id","count"),\n'
'    delivery_rate=("delivered","mean")\n'
').assign(pct=lambda x: x["revenue"]/total_rev*100)\n'
'biz_rev = df.groupby("BusinessId")["Price"].sum().sort_values(ascending=False)')

data_box(doc, [
    'סה"כ הכנסות: ₪3,407,690  |  סה"כ הזמנות: 28,645  |  AOV כלל: ₪118.96',
    '',
    'פירוק לפי אזור (מיון לפי הכנסה):',
    '  Umm al-Fahem:       ₪1,037,006 (30.4%) | 8,784 הזמנות | AOV ₪118.06 | מסירה 58.6%',
    '  Sakhnin:             ₪641,057 (18.8%) | 4,569 הזמנות | AOV ₪140.31 | מסירה 75.7%',
    '  Kfar Qaree - Arara:  ₪429,095 (12.6%) | 3,709 הזמנות | AOV ₪115.69 | מסירה 55.4%',
    '  Nazareth:            ₪372,434 (10.9%) | 2,948 הזמנות | AOV ₪126.33 | מסירה 74.5%',
    '  Baqa al-Gharbiyye:   ₪360,947 (10.6%) | 3,199 הזמנות | AOV ₪112.83 | מסירה 57.0%',
    '  Tamra - Kabul:       ₪237,653  (7.0%) | 2,104 הזמנות | AOV ₪112.95 | מסירה 65.4%',
    '  Shefa-Amr:           ₪170,031  (5.0%) | 1,493 הזמנות | AOV ₪113.89 | מסירה 58.3%',
    '  Lod:                  ₪58,188  (1.7%) |   637 הזמנות | AOV  ₪91.35 | מסירה 94.8%',
    '  Jisr az-Zarqa:        ₪42,335  (1.2%) |   420 הזמנות | AOV ₪100.80 | מסירה 96.2%',
    '  Kafr Qasem:           ₪33,268  (1.0%) |   461 הזמנות | AOV  ₪72.16 | מסירה 97.0%',
    '  Ramla:                ₪25,676  (0.8%) |   321 הזמנות | AOV  ₪79.99 | מסירה 91.6%',
    '',
    'פירוט עסקים: 1,274 עסקים פעילים',
    '  TOP 3: עסק 3 (₪95,058 | 2.79%) | עסק 4 (₪94,401 | 2.77%) | עסק 173 (₪54,929 | 1.61%)',
    '  TOP 3 יחד: ₪244,388 = 7.2% מסה"כ ההכנסות',
    '  TOP 10 יחד: ₪481,970 = 14.1% מסה"כ ההכנסות',
])

answer(doc, '5 האזורים הראשונים (Umm al-Fahem, Sakhnin, Kfar Qaree-Arara, Nazareth, Baqa al-Gharbiyye) מהווים 83.3% מסה"כ ההכנסות — ריכוזיות גבוהה.')
answer_bullet(doc, 'Umm al-Fahem בולטת: 30.4% מהכנסות הפלטפורמה כולה. ירידה בביצועים שם תורגש מיד ברמת החברה.')
answer_bullet(doc, 'Sakhnin מייצרת AOV גבוה יחסית (₪140.31) — 18% יותר מהממוצע. ייתכן שם mix-product שונה (הזמנות גדולות יותר).')
answer_bullet(doc, 'ריכוזיות עסקית נמוכה: TOP 3 עסקים = 7.2% בלבד. הפלטפורמה אינה תלויה בעסק יחיד — זה טוב לחוסן.')
add_chart(doc, 'q5a_revenue_area.png', caption='תרשים 7 — הכנסות לפי אזור (עמודות) מול שיעור מסירה (נקודות כחולות)')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ב | הקורלציה בין אזורי הכנסה גבוהה לאיכות מסירה')

data_box(doc, [
    'מתאם פירסון בין הכנסה לשיעור מסירה: r = -0.623 (קורלציה שלילית!)',
    '',
    'TOP 5 לפי הכנסה:  Umm al-Fahem | Sakhnin | Kfar Qaree-Arara | Nazareth | Baqa al-Gharbiyye',
    'TOP 5 לפי מסירה:  Kafr Qasem | Jisr az-Zarqa | Lod | Ramla | Sakhnin',
    '',
    'חפיפה בין שתי הרשימות: רק Sakhnin מופיעה בשתיהן!',
    '',
    'הפרדוקס — ערים גדולות (הכנסה גבוהה) = מסירה גרועה:',
    '  Umm al-Fahem (הכנסה #1):        מסירה 58.6% — פחות מ-6 מ-10 הזמנות מגיעות',
    '  Kfar Qaree - Arara (הכנסה #3):  מסירה 55.4% — גרוע ביותר',
    '  Baqa al-Gharbiyye (הכנסה #5):   מסירה 57.0%',
    '',
    'ערים קטנות (הכנסה נמוכה) = מסירה מצוינת:',
    '  Kafr Qasem (הכנסה #10):   מסירה 97.0%',
    '  Jisr az-Zarqa (הכנסה #9): מסירה 96.2%',
    '  Lod (הכנסה #8):           מסירה 94.8%',
])

answer(doc, 'קיים פרדוקס מובהק: האזורים שמייצרים את מרבית ההכנסה הם בדיוק אלה שמתפקדים הכי גרוע במסירה. מקדם הקורלציה של -0.623 מאשר שהקשר הפוך וחזק.')
answer_bullet(doc, 'ההסבר הסביר: ערים גדולות = ביקוש גבוה = מחסור בנהגים. אותה בעיית אספקה שזיהינו בשאלה 3 מתבטאת כאן גם בממד הפיננסי.')
answer_bullet(doc, 'ה"אובדן" הנסתר: אם Umm al-Fahem הייתה מגיעה ל-90% מסירה (במקום 58.6%) — הפלטפורמה הייתה מרוויחה ~31% יותר הכנסות ממנה.')
warning_box(doc, 'הפלטפורמה גובה תשלום על הזמנות שלא מגיעות? אם לא — שיעור מסירה של 58.6% ב-Umm al-Fahem אומר שכמעט ₪430,000 (41.4% מ-₪1,037,006) לא הומרו להכנסה מנוהלת. בעיה חמורה שדורשת בדיקת מדיניות ביטולים.')
add_chart(doc, 'q5b_scatter.png', caption='תרשים 8 — קורלציה שלילית בין הכנסות לשיעור מסירה לפי אזור (r = -0.62)')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ג | תלות ב-TOP 3 עסקים')

data_box(doc, [
    'TOP 3 עסקים (3, 4, 173): ₪244,388 = 7.2% מסה"כ הכנסות',
    'אם TOP 3 עוצרים — הפלטפורמה מאבדת 7.2% מהכנסתה',
    '',
    'TOP 10 עסקים: ₪481,970 = 14.1% מסה"כ הכנסות',
    'כלומר: 1,264 העסקים הנותרים מייצרים 85.9% מההכנסות',
    '',
    'ההכנסה מפוזרת בין 1,274 עסקים — פיזור טוב, אין תלות אחד:',
    '  עסק הגדול ביותר (עסק 3): 2.79% מסה"כ ההכנסות בלבד',
    '  פי שמינית מהתלות על Umm al-Fahem (30.4%) כאזור',
])

answer(doc, 'HAAT אינה תלויה ב-TOP 3 עסקים — אם הם יעצרו פעילות, הפלטפורמה תאבד 7.2% בלבד מהכנסותיה. זהו מצב בריא מבחינת גיוון תיק העסקים.')
answer_bullet(doc, 'הריכוזיות האמיתית נמצאת ברמת האזורים — 5 ערים = 83% מהכנסות. כאן הסיכון גבוה יותר.')
answer_bullet(doc, 'עסק 3 ועסק 4 (AOV גבוה) — כדאי לבדוק אם מדובר בעסק קטרינג/חלוקה, שכן נפח הזמנות שלהם (906, 650) גבוה מאוד ביחס לאחרים.')
add_chart(doc, 'q5c_top_businesses.png', caption='תרשים 9 — TOP 10 עסקים לפי הכנסה | TOP 3 (אדום) = 7.2% מסה"כ')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ד | המלצה לצוות המסחרי')

answer(doc, 'על בסיס הניתוח, אלה 3 העדיפויות לצוות המסחרי:')
answer_bullet(doc,
    'עדיפות 1 — שיפור כיסוי נהגים בערים הגדולות: '
    'כל אחוז שיפור במסירה ב-Umm al-Fahem שווה ~₪10,000 הכנסה נוספת (בהינתן AOV ₪118). '
    'זה ה-ROI הגבוה ביותר האפשרי. '
    'המלצה: גיוס/תמריצים לנהגים בשעות ערב (17:00–22:00) דווקא ב-5 הערים הגדולות.',
    color=DARK_BLUE)
answer_bullet(doc,
    'עדיפות 2 — צמיחה בAOV של Sakhnin: '
    'Sakhnin היא הנוסחה המנצחת — AOV ₪140 (גבוה ב-18%) + שיעור מסירה 75.7% (שלישי). '
    'כדאי לבדוק מה גורם ל-AOV גבוה שם (סוגי עסקים? גודל הזמנות?) ולשכפל זאת לאזורים אחרים.',
    color=DARK_BLUE)
answer_bullet(doc,
    'עדיפות 3 — גיוס עסקים בלוד ו-Ramla: '
    'שיעורי מסירה של 91-95% עם נפח נמוך יחסית = ביקוש לא מסופק. '
    'גיוס עסקים חדשים שם יניב הכנסה "זולה" — הצי הקיים מסוגל לטפל בנפח גדול יותר.',
    color=DARK_BLUE)

insight(doc, 'הממצא הבסיסי: הפלטפורמה "משאירה כסף על השולחן" בערים הגדולות בגלל כשלי מסירה. תיקון בעיית הנהגים יניב עלייה ישירה בהכנסות ללא צורך בגיוס לקוחות חדשים.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  חלק ג — חשיבה עסקית
# ══════════════════════════════════════════════════════════════════════════════
part_banner(doc, 'חלק ג — חשיבה עסקית')

# ════════════════════════════════════════════════════════════════
#  שאלה 6 — שימור לקוחות (קאסטומר ריטנשן)
# ════════════════════════════════════════════════════════════════
q_header(doc, 'שאלה 6', 'שימור לקוחות (קאסטומר ריטנשן)')
p(doc, 'הנחיה: נתח חזרת לקוחות לפי אזורים ולפי קוהורט זמן.',
  italic=True, size=10, color=MID_GREY, sb=0, sa=6)

sub(doc, 'א | אחוז הזמנות מלקוחות חוזרים — לפי אזור')

code_block(doc,
'df_sorted = df.sort_values("OrderDate")\n'
'df_sorted["is_returning"] = df_sorted.duplicated(subset="UserId", keep="first")\n'
'area_ret  = df_sorted.groupby("AreaName").agg(\n'
'    orders=("Id","count"), returning=("is_returning","sum")\n'
').assign(ret_rate=lambda x: x["returning"]/x["orders"]*100)')

data_box(doc, [
    'שיעור הזמנות מלקוחות חוזרים — כלל הפלטפורמה: 22.8%',
    '(כלומר: 77.2% מהזמנות הן מלקוחות שמזמינים בפעם הראשונה)',
    '',
    'לפי אזור (מיון יורד לפי שיעור חזרה):',
    '  Kafr Qasem:        34.9% (161/461 הזמנות חוזרות)   ← מוביל',
    '  Ramla:             29.3% (94/321 הזמנות חוזרות)',
    '  Umm al-Fahem:      27.1% (2,384/8,784 הזמנות חוזרות)',
    '  Lod:               26.5% (169/637 הזמנות חוזרות)',
    '  Kfar Qaree - Arara: 23.6% (877/3,709 הזמנות חוזרות)',
    '  Baqa al-Gharbiyye: 21.0% (671/3,199 הזמנות חוזרות)',
    '  Shefa-Amr:         20.8% (310/1,493 הזמנות חוזרות)',
    '  Jisr az-Zarqa:     19.8% (83/420 הזמנות חוזרות)',
    '  Tamra - Kabul:     19.3% (406/2,104 הזמנות חוזרות)',
    '  Nazareth:          18.6% (547/2,948 הזמנות חוזרות)',
    '  Sakhnin:           18.2% (833/4,569 הזמנות חוזרות)  ← גרוע ביותר',
])

answer(doc, 'בממוצע, 22.8% מההזמנות מגיעות מלקוחות שכבר הזמינו לפחות פעם אחת קודם. זהו שיעור נמוך יחסית לפלטפורמת משלוחים בוגרת.')
answer_bullet(doc, 'Kafr Qasem מובילה ב-34.9% — כמעט 1 מכל 3 הזמנות שם היא מלקוח חוזר. ייתכן שם קהילה קטנה ולכידה עם פחות תחרות.')
answer_bullet(doc, 'Sakhnin גרועה ביותר ב-18.2% — מפתיע, כי שיעור המסירה שם סביר (75.7%). ייתכן שמדובר בבעיית מוצר (תמחור, מגוון עסקים) ולא בבעיית שירות.')
answer_bullet(doc, 'Umm al-Fahem — למרות שיעור מסירה גרוע (58.6%), שיעור החזרה (27.1%) גבוה מהממוצע. כנראה שאין תחרות וללקוחות אין חלופה.')
add_chart(doc, 'q6a_returning_area.png', caption='תרשים 10 — שיעור הזמנות מלקוחות חוזרים לפי אזור')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ב | קוהורט מרץ 2024 — חזרה תוך 30 ו-60 יום')

code_block(doc,
'march_new = df_sorted[\n'
'    (df_sorted["is_first"]) & (df_sorted["Month"] == "2024-03")\n'
']['
'UserId"].unique()\n'
'cohort = df[df["UserId"].isin(march_new)].copy()\n'
'cohort = cohort.merge(first_orders, on="UserId")\n'
'cohort["days"] = (cohort["OrderDate"] - cohort["first_date"]).dt.days\n'
'repeat_30 = cohort[(cohort["days"]>0) & (cohort["days"]<=30)]["UserId"].nunique()\n'
'repeat_60 = cohort[(cohort["days"]>0) & (cohort["days"]<=60)]["UserId"].nunique()')

data_box(doc, [
    'קוהורט: לקוחות שהזמנתם הראשונה היתה במרץ 2024',
    'גודל הקוהורט: 4,941 לקוחות חדשים',
    '',
    'תוך 30 יום  → הזמינו שוב: 970 לקוחות = 19.6%',
    'תוך 60 יום  → הזמינו שוב: 1,367 לקוחות = 27.7%',
    '',
    'פרשנות:',
    '  ~1 מכל 5 לקוחות חדשים חוזר תוך חודש',
    '  ~1 מכל 4 לקוחות חדשים חוזר תוך חודשיים',
    '  שיעור גדילה בין 30 ל-60 יום: +8.1 נקודות אחוז (+41% יחסי)',
    '  לקוחות רבים שחוזרים עושים זאת בחלון 31-60 יום (397 נוספים)',
])

answer(doc, 'מתוך 4,941 לקוחות חדשים במרץ 2024 — רק 19.6% חזרו תוך 30 יום, ו-27.7% תוך 60 יום.')
answer_bullet(doc, 'שיעור 19.6% ל-30 יום נמוך לפלטפורמת פוד-דליוורי. פלטפורמות בשלות מגיעות ל-30-40% ריטנשן ב-30 יום.')
answer_bullet(doc, 'הקפיצה הגדולה בין 30 ל-60 יום (8.1%) מצביעה שלקוחות רבים מחליטים לחזור אחרי "מנוחה" של חודש — ייתכן שזה קשור לפרקי תשלום (משכורות חודשיות) או לאירועים חברתיים.')
answer_bullet(doc, 'חלון קריטי: 72.3% מהלקוחות שלא חזרו תוך 60 יום — כנראה לא יחזרו. זהו החלון לפעולת שימור אקטיבית.')
insight(doc, 'הקמפיין האידיאלי: לקוחות שלא הזמינו 15 יום לאחר הזמנה ראשונה — לשלוח להם קוד הנחה. החלון הכי קריטי לשימור הוא יום 10-20 לאחר ההזמנה הראשונה.')
add_chart(doc, 'q6b_retention.png', caption='תרשים 11 — קוהורט מרץ 2024 (עוגה) ושיעור שימור 30 יום לפי אזור (עמודות)')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ג | השוואת שיעור שימור 30 יום לפי אזור')

data_box(doc, [
    'שיעור שימור 30 יום (לקוחות שהזמינו שוב תוך 30 יום מהזמנתם הראשונה):',
    '',
    '  Ramla:              20.3% (46/227 לקוחות)    ← מוביל',
    '  Lod:                15.4% (72/468 לקוחות)',
    '  Kafr Qasem:         15.3% (46/300 לקוחות)',
    '  Jisr az-Zarqa:      13.9% (47/337 לקוחות)',
    '  Umm al-Fahem:       13.8% (885/6,400 לקוחות)',
    '  Kfar Qaree - Arara: 12.3% (347/2,832 לקוחות)',
    '  Shefa-Amr:          11.9% (141/1,183 לקוחות)',
    '  Baqa al-Gharbiyye:  10.8% (272/2,528 לקוחות)',
    '  Tamra - Kabul:      10.5% (178/1,698 לקוחות)',
    '  Nazareth:           10.1% (243/2,401 לקוחות)',
    '  Sakhnin:             9.7% (361/3,736 לקוחות)  ← גרוע ביותר',
    '',
    'פער בין הטוב לגרוע: 20.3% (Ramla) vs 9.7% (Sakhnin) = פי 2.1',
])

answer(doc, 'Ramla מובילה עם שיעור שימור 30 יום של 20.3% — כמעט פי 2 מ-Sakhnin שנמצאת בתחתית (9.7%).')
answer_bullet(doc, 'מעניין: Ramla מובילה בשימור (20.3%) למרות שאינה מובילה בשיעור חזרה כללי (29.3% כלל-תקופה). הסבר: ייתכן שלקוחות שם חוזרים מהר יחסית.')
answer_bullet(doc, 'Sakhnin גרועה גם בשיעור חזרה (18.2%) וגם בשימור 30 יום (9.7%) — סימן שהבעיה עמוקה יותר מאשר רק שיוך נהגים.')
answer_bullet(doc, 'ערים עם שיעור מסירה גבוה (Lod, Kafr Qasem) מציגות שימור טוב — קשר ישיר: מסירה טובה = לקוח מרוצה = חוזר.')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ד | המלצה לשיפור שימור ב-Sakhnin (האזור הגרוע ביותר)')

answer(doc, 'Sakhnin מציגה שיעור שימור 30 יום של 9.7% בלבד — הגרוע מכולם. להלן תוכנית חקירה וטיפול:')
answer_bullet(doc,
    'מה לחקור: '
    'ראשית, יש לבדוק האם בעיית השימור היא בחוויית המשלוח (זמן המתנה ארוך? '
    'כשלי מסירה ספציפיים?) או במוצר (מגוון עסקים מצומצם? תמחור גבוה ביחס לחלופות?). '
    'נוסף על כך יש לבדוק האם קיימת תחרות מקומית חזקה (סופרמרקטים, '
    'מסעדות עם משלוח עצמאי) שמושכת לקוחות חזרה.',
    color=DARK_BLUE)
answer_bullet(doc,
    'אילו נתונים נדרשים: '
    'דירוגי לקוחות לפי הזמנה (לזהות אם חוויית הזמנה ראשונה היתה שלילית), '
    'נתוני ביטולים ב-Sakhnin, השוואת מגוון עסקים פעילים לעומת ערים עם שימור גבוה, '
    'וסקר אופציונלי מלקוחות שלא חזרו.',
    color=DARK_BLUE)
answer_bullet(doc,
    'מה לבדוק (א/ב טסטינג): '
    'ניסוי 1 — קוד הנחה של 15% ללקוחות ב-Sakhnin שלא הזמינו 14 יום לאחר הזמנה ראשונה, '
    'לעומת קבוצת ביקורת ללא הנחה. '
    'ניסוי 2 — גיוס 2-3 עסקים חדשים ב-Sakhnin (מסוגים שחסרים) וניתוח '
    'השפעת הרחבת הקטלוג על שיעורי החזרה. '
    'מדד הצלחה: שיפור ב-3+ נקודות אחוז בשימור 30 יום תוך רבעון.',
    color=DARK_BLUE)

insight(doc, 'Sakhnin היא פרדוקס: הכנסות גבוהות (#2), AOV גבוה (₪140), שיעור מסירה סביר (75.7%) — אבל שימור גרוע. זה מצביע על בעיה במוצר או בתחרות, לא בתפעול. גישת הטיפול צריכה להיות שיווקית ומוצרית, לא תפעולית.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  שאלה 7 — דפוסי הזמנות לאורך זמן
# ════════════════════════════════════════════════════════════════
q_header(doc, 'שאלה 7', 'דפוסי הזמנות לאורך זמן (אורדר פטרנז אובר טיים)')
p(doc, 'הנחיה: נתח כיצד התנהגות ההזמנות משתנה לאורך תקופת הדאטהסט. תאר דפוסים, שינויים ואנומליות.',
  italic=True, size=10, color=MID_GREY, sb=0, sa=6)

sub(doc, 'א | אנומליה 1 — נפילת שבוע 11–17 מרץ: −40.1%')

code_block(doc,
'df["Week"] = df["OrderDate"].dt.to_period("W")\n'
'weekly = df.groupby("Week").agg(orders=("Id","count"), delivery_rate=("delivered","mean"))\n'
'weekly["wow_pct"] = weekly["orders"].pct_change() * 100')

data_box(doc, [
    'שבוע 04–10 מרץ:  1,894 הזמנות | AOV ₪124.01 | מסירה 69.6%',
    'שבוע 11–17 מרץ:  1,134 הזמנות | AOV ₪106.73 | מסירה 72.2%  ← נפילה חדה!',
    'שינוי: −760 הזמנות = −40.1% בשבוע אחד',
    'שבוע 18–24 מרץ:  1,575 הזמנות | מסירה 74.8%  ← התאוששות חלקית +38.9%',
    '',
    'מיוחד: ב-4 הימים 11–14 מרץ נפח יומי ירד ל-121–129 הזמנות (שפל כלל-תקופה)',
    'ובאותה עת — שיעור המסירה דווקא עלה ל-66–73%',
])

answer(doc, 'הנפילה הדרמטית החלה ב-11 מרץ 2024 — תאריך שמתואם במדויק לתחילת חודש הרמדאן 2024.')
answer_bullet(doc, 'השערה: ב-11 מרץ 2024 החל הרמדאן. בקהילה הערבית (שהיא בסיס הלקוחות של HAAT) השינוי בדפוסי האכילה מוביל לירידה בהזמנות בשעות היום. בשעות הלילה הנפח חוזר — אך ה"שבוע" הסטטיסטי נראה נמוך יותר.')
answer_bullet(doc, 'ראיה תומכת: ממש בשבוע 11–17 מרץ שיעור המסירה עלה מ-69.6% ל-72.2% — כי הנהגים שזמינים בשעות השפל (בוקר/צהריים) מסוגלים לכסות את ההזמנות הבודדות בקלות.')
answer_bullet(doc, 'השלכה עסקית: HAAT חשופה מאוד לרמדאן ולחגים מוסלמיים. ירידה של 40% בנפח שבועי גורמת לנהגים להתפזר ולהכנסות לרדת. מנגד — אחרי פטאר (שבירת הצום בערב) יש עלייה חדה בביקוש שצריך להיות מכוסה בנהגים.')
answer_bullet(doc, 'נתונים להוכחה: תאריכי הרמדאן ב-2024 (11 מרץ – 9 אפריל), פירוט נפח לפי שעה בשבוע 11–17 מרץ (האם השיא עבר מ-20:00 לשעות מאוחרות יותר?).')
insight(doc, 'ממצא אסטרטגי: אם HAAT מוכנה מראש לרמדאן עם תמחור דינמי, מבצעי שבירת צום ווורקשיפטים מותאמים לנהגים — אפשר להפוך את העונה הזו מנפילה להזדמנות.')

add_chart(doc, 'q7a_weekly_trend.png', caption='תרשים 12 — מגמה שבועית: נפח הזמנות, שיעור מסירה ו-AOV | האזורים האדומים = שבועות אנומליה')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ב | אנומליה 2 — שבוע 8–14 אפריל: שיעור מסירה שפל של 48.4%')

data_box(doc, [
    'שבוע 01–07 אפריל: 1,363 הזמנות | מסירה 62.1%',
    'שבוע 08–14 אפריל: 1,573 הזמנות | מסירה 48.4%  ← שפל כלל-תקופה!',
    'שבוע 15–21 אפריל: 1,987 הזמנות | מסירה 67.4%  ← קפיצה חזרה',
    '',
    'AOV באותו שבוע: ₪128.86 — הגבוה ביותר בכל התקופה',
    'הנפח דווקא עלה (+15.4%) אבל המסירה התרסקה',
])

answer(doc, 'שבוע 8–14 אפריל הוא הממצא החריג ביותר בכל הנתונים: נפח עלה ב-15% אך שיעור המסירה צנח ל-48.4% — ממוצע מסירה של פחות מהזמנה אחת מתוך שתיים.')
answer_bullet(doc, 'השערה 1 — אירוע תפעולי: כשל טכני, מחלה המונית של נהגים, או אירוע ביטחוני שגרם לנהגים להימנע מיציאה. 8–14 אפריל 2024 כולל את חג הפסח (פסח 2024: 22–30 אפריל — לא מדויק) וסמוך לתקיפת איראן (13 אפריל 2024!).')
answer_bullet(doc, 'השערה 2 — התקפת איראן: ב-13–14 אפריל 2024 ישראל חוותה תקיפת טילים ומל"טים. אם חלק מנהגי HAAT הפסיקו לעבוד בסוף השבוע הזה, זה מסביר את שפל המסירה.')
answer_bullet(doc, 'ראיה תומכת: AOV גבוה במיוחד (₪128.86) — ייתכן שרק הזמנות גדולות/נחוצות הוזמנו, ולקוחות בחרו מוצרים חיוניים.')
answer_bullet(doc, 'נתונים להוכחה: כמות נהגים פעילים ב-13–14 אפריל לעומת שאר ימי השבוע, וגיאוגרפיה של כשלי המסירה.')
warning_box(doc, 'שבוע 8–14 אפריל: 48.4% מסירה = כ-840 הזמנות שלא נמסרו. ב-AOV של ₪128.86, אלה כ-₪108,000 בפוטנציאל הכנסה שאבד — ובלקוחות מתוסכלים.')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ג | דפוס יום בשבוע — שבת מוביל, ראשון הכי יעיל')

data_box(doc, [
    'ממוצע הזמנות ליום לפי יום:',
    '  שבת:    271.1 הזמנות/יום | AOV ₪122.76 | מסירה 63.1%  ← שיא נפח',
    '  חמישי:  259.1 הזמנות/יום | AOV ₪120.70 | מסירה 64.8%',
    '  שישי:   249.1 הזמנות/יום | AOV ₪120.59 | מסירה 64.3%',
    '  ראשון:  227.3 הזמנות/יום | AOV ₪118.14 | מסירה 67.5%  ← מסירה הטובה ביותר',
    '  שני:    216.2 הזמנות/יום | AOV ₪117.33 | מסירה 66.0%',
    '  שלישי:  220.0 הזמנות/יום | AOV ₪114.67 | מסירה 65.9%',
    '  רביעי:  214.3 הזמנות/יום | AOV ₪117.11 | מסירה 65.3%',
    '',
    'פיק שעה: פברואר — 20:00 | מרץ–מאי — 21:00 (הזזה של שעה)',
])

answer(doc, 'שבת היא יום השיא בנפח (271 הזמנות בממוצע) — 26% יותר מימי שני-רביעי. חמישי ושישי (ערב סוף שבוע) גם הם גבוהים.')
answer_bullet(doc, 'ראשון הוא הפרדוקס: נפח בינוני (227) אך מסירה הכי טובה (67.5%). כנראה הנהגים שסיימו את שבת ממשיכים לעבוד ביום ראשון.')
answer_bullet(doc, 'שבת: נפח שיא אבל מסירה גרועה (63.1%) — עומס על הצי. זו הזדמנות ברורה לגיוס נהגים ספציפית לשבת.')
answer_bullet(doc, 'הזזת שעת השיא מ-20:00 לפברואר ל-21:00 ממרץ ואילך — ייתכן שקשורה לרמדאן (פטאר בשעה מאוחרת יותר) ולהארכת שעות הפעילות בעונת הקיץ.')
insight(doc, 'תובנה תפעולית: גיוס נהגים ממוקד לשבת (17:00–23:00) ולחמישי-שישי יניב את ה-ROI הגבוה ביותר — הביקוש כבר שם, חסרים הנהגים.')

add_chart(doc, 'q7b_dow.png', caption='תרשים 13 — ממוצע הזמנות, שיעור מסירה ו-AOV לפי יום בשבוע')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  שאלה 8 — בדיקת איכות נתונים
# ════════════════════════════════════════════════════════════════
q_header(doc, 'שאלה 8', 'בדיקת איכות נתונים (דאטה קוואליטי צ\'ק)')
p(doc, 'הנחיה: חקור את טבלת payment_methods_user. בדוק האם הקשרים בין לקוחות לאמצעי תשלום נראים תקינים.',
  italic=True, size=10, color=MID_GREY, sb=0, sa=6)

sub(doc, 'א | מבנה הטבלה וממצאים ראשוניים')

code_block(doc,
'pm = pd.read_excel("HAAT_DA_Dataset.xlsx", sheet_name="payment_methods_user")\n'
'print(pm.columns.tolist())   # ["Id", "UserId", "Token"]\n'
'print(f"שורות: {len(pm)} | UserId ייחודיים: {pm[\'UserId\'].nunique()}")\n'
'tpu = pm.groupby("UserId")["Token"].count().sort_values(ascending=False)')

data_box(doc, [
    'מבנה הטבלה: Id (מזהה שורה) | UserId (מזהה לקוח) | Token (טוקן כרטיס אשראי)',
    'סה"כ שורות: 50,174  |  UserId ייחודיים: 15,287',
    '',
    'פיזור טוקנים לכל לקוח:',
    '  1 טוקן: 4,643 לקוחות (30.4%)  — הכי שכיח',
    '  2 טוקנים: 3,438 לקוחות (22.5%)',
    '  3 טוקנים: 2,407 לקוחות (15.7%)',
    '  5+ טוקנים: 3,211 לקוחות (21.0%)',
    '  10+ טוקנים: 571 לקוחות (3.7%)',
    '',
    'ערכים חסרים: 0 בכל עמודה  |  פורמט Token: TOK_XXXXXXXX (12 תווים, אחיד)',
])

answer(doc, 'הטבלה מייצגת טוקניזציה של כרטיסי אשראי — כל Token הוא ייצוג מוצפן של כרטיס ספציפי. לכן לקוח יכול להחזיק מספר טוקנים (מספר כרטיסים).')
answer_bullet(doc, 'הפורמט אחיד לחלוטין: כל הטוקנים מתחילים ב-TOK_ ובאורך 12 תווים — אין ערכים פגומים בפורמט.')
answer_bullet(doc, 'אין ערכים חסרים — הטבלה נקייה ברמת השדות.')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ב | ממצא חריג 1 — לקוחות עם מאות טוקנים')

data_box(doc, [
    'TOP 5 לקוחות לפי מספר טוקנים:',
    '  UserId 11344: 161 טוקנים | 2 הזמנות בלבד | 1 מזומן',
    '  UserId 20614: 160 טוקנים | 2 הזמנות בלבד | 0 מזומן',
    '  UserId 10719: 107 טוקנים | 1 הזמנה בלבד  | 0 מזומן',
    '  UserId 19970:  81 טוקנים | 3 הזמנות בלבד | 0 מזומן',
    '  UserId 12332:  68 טוקנים | 1 הזמנה בלבד  | 0 מזומן',
    '',
    'לקוחות עם 10+ טוקנים: 571 לקוחות',
])

warning_box(doc,
    'UserId 11344 רשם 161 כרטיסי אשראי שונים — אך הזמין רק פעמיים. '
    'זוהי דגל אדום קלאסי לאחת מהאפשרויות הבאות: '
    '(1) ניסיון הונאה — בדיקת תוקף כרטיסים גנובים על פלטפורמת HAAT, '
    '(2) באג במערכת שרשם טוקנים כפולים לאותו לקוח, '
    '(3) חשבון טסט של מפתח שרץ על הסביבה הפרודקשן.')

answer(doc, 'לקוחות עם 50+ טוקנים ופחות מ-5 הזמנות הם ממצא חשוד מאוד שמחייב בדיקה מיידית.')
answer_bullet(doc, 'הצעד הבא: להצליב את UserId 11344 ו-20614 מול לוגי המערכת — האם כל 161 הטוקנים נרשמו בפרק זמן קצר? מאותה כתובת IP? זה יעיד על ניסיון Card-Testing.')
answer_bullet(doc, 'המלצה: להוסיף ריט לימיט (Rate Limit) על הוספת כרטיסים — לדוגמה מקסימום 5 כרטיסים ב-24 שעות לכל חשבון.')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ג | ממצא חריג 2 — 1,258 טוקנים כפולים')

data_box(doc, [
    'טוקנים כפולים (אותו Token בשורות שונות): 1,258',
    'טוקנים ייחודיים: 48,916 מתוך 50,174 שורות (97.5% ייחודיות)',
    '',
    'כל הטוקנים שייכים לאותו UserId? (בדיקה):',
    '  אם כן — סתם רישום כפול (באג בלוגיק הכתיבה לDB)',
    '  אם לא — כרטיס אחד משויך לשני לקוחות שונים (חמור יותר)',
])

answer(doc, '1,258 טוקנים מופיעים יותר מפעם אחת. Token אמור לייצג כרטיס ייחודי — כפילויות מעידות על באג ברישום.')
answer_bullet(doc, 'הצעד הבא: לבדוק אם הכפילויות הן לאותו UserId (באג בלוגיק "הוסף כרטיס") או לUserId שונים (פרצת אבטחה — שיתוף כרטיס בין חשבונות).')
answer_bullet(doc, 'המלצה: להוסיף UNIQUE constraint על עמודת Token בDB כדי למנוע רישום כפול עתידי.')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ד | ממצא חריג 3 — קרוס רפרנס עם טבלת הזמנות')

data_box(doc, [
    'לקוחות ייחודיים בהזמנות:        22,110',
    'לקוחות ייחודיים ב-payment:       15,287',
    'לקוחות בשתי הטבלאות:            15,280',
    'רק ב-payment (לא הזמינו):            7  ← סביר',
    'רק בהזמנות (אין טוקן):           6,830  ← צפוי (משלמי מזומן)',
    '',
    'בדיקת עקביות תשלום:',
    '  משלמי מזומן (PaymentMethod=0) עם טוקן:    9,190 / 15,945 (57.6%)',
    '  משלמי אשראי (PaymentMethod=1) ללא טוקן:      78 / 7,041  (1.1%)  ← חריג!',
    '  לקוחות שמשלמים גם מזומן וגם אשראי:          876 לקוחות',
])

answer(doc, 'הממצא הקריטי: 78 לקוחות שיצאה מהם הזמנה עם PaymentMethod=1 (אשראי) — אבל הם אינם מופיעים בטבלת payment_methods_user כלל.')
answer_bullet(doc, 'הסבר אפשרי: תשלום דרך ספק חיצוני (PayPal, Apple Pay) שאינו מייצר טוקן פנימי, או מחיקת הכרטיס לאחר ההזמנה.')
answer_bullet(doc, 'ממצא תקין: 6,830 משלמי מזומן ללא טוקן — לגמרי הגיוני, מזומן לא דורש רישום כרטיס.')
answer_bullet(doc, 'ממצא מעניין: 876 לקוחות משלמים גם מזומן וגם אשראי — לקוחות גמישים שיש לתמרץ לאשראי (מעקב טוב יותר, אפשרות לתכנית נאמנות).')
insight(doc, 'הנתונים בסך הכל נקיים ועקביים — אין שחיתות מסיבית. 3 הממצאים (טוקנים מרובים, כפילויות, 78 חריגים) הם בעיות ידועות בכל מערכת תשלומים. הטיפול הקריטי ביותר הוא UserId 11344/20614 בשל חשד ל-Card Testing.')

add_chart(doc, 'q8_data_quality.png', caption='תרשים 14 — פיזור טוקנים ללקוח (שמאל) והשוואת טבלאות (ימין)')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  שאלה 10 — בונוס: חישוב בונוסים שבועיים לשליחים
# ══════════════════════════════════════════════════════════════════════════════
q_header(doc, 'שאלה 10 — בונוס', 'חישוב בונוסים שבועיים לשליחים (Courier Bonus Calculation)')
p(doc, 'הנחיה: כתוב פונקציית Python המקבלת את הדאטאפריים ותאריכי שבוע ומחזירה טבלת בונוסים לשליחים לפי חוקי אזור.',
  italic=True, size=10, color=MID_GREY, sb=0, sa=6)

sub(doc, 'א | חוקי הבונוס')

data_box(doc, [
    'AreaId 1  — סוף שבוע (שישי + שבת) בלבד: כל 50 מסירות → 100 ₪ בונוס',
    'AreaId 4  — כל ימות השבוע: כל מסירה מעל 300 ← 3 ₪ למסירה',
    'AreaId 10 — סוף שבוע (שישי + שבת) בלבד: כל 50 מסירות → 150 ₪ בונוס',
    '',
    'הגדרות:',
    '  • Delivered = ArriveDate is not NULL',
    '  • Weekend   = Friday (DayOfWeek=4) or Saturday (DayOfWeek=5)',
    '  • פלט: DriverId, AreaId, TotalDelivered, BonusAmount',
])

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ב | הפונקציה')

code_block(doc,
'def calc_weekly_bonus(orders_df, week_start, week_end):\n'
'    """\n'
'    מחשבת בונוסים שבועיים לשליחים לפי חוקי אזור.\n'
'    """\n'
'    ws = pd.Timestamp(week_start)\n'
'    we = pd.Timestamp(week_end) + pd.Timedelta(days=1) - pd.Timedelta(seconds=1)\n'
'\n'
'    # סינון לשבוע + מסירות בלבד + נהג ידוע\n'
'    w = orders_df[\n'
'        (orders_df["OrderDate"] >= ws) &\n'
'        (orders_df["OrderDate"] <= we) &\n'
'        (orders_df["delivered"] == True) &\n'
'        (orders_df["DriverId"].notna())\n'
'    ].copy()\n'
'\n'
'    results = []\n'
'\n'
'    # AreaId 1 — סוף שבוע, 100 ₪ לכל 50 מסירות\n'
'    a1 = w[(w["AreaId"] == 1) & (w["IsWeekend"] == True)]\n'
'    g1 = a1.groupby("DriverId").size().reset_index(name="TotalDelivered")\n'
'    g1["AreaId"]      = 1\n'
'    g1["BonusAmount"] = (g1["TotalDelivered"] // 50) * 100\n'
'    results.append(g1)\n'
'\n'
'    # AreaId 4 — כל הימים, 3 ₪ מעל מסירה 300\n'
'    a4 = w[w["AreaId"] == 4]\n'
'    g4 = a4.groupby("DriverId").size().reset_index(name="TotalDelivered")\n'
'    g4["AreaId"]      = 4\n'
'    g4["BonusAmount"] = g4["TotalDelivered"].apply(lambda n: max(0, n - 300) * 3)\n'
'    results.append(g4)\n'
'\n'
'    # AreaId 10 — סוף שבוע, 150 ₪ לכל 50 מסירות\n'
'    a10 = w[(w["AreaId"] == 10) & (w["IsWeekend"] == True)]\n'
'    g10 = a10.groupby("DriverId").size().reset_index(name="TotalDelivered")\n'
'    g10["AreaId"]      = 10\n'
'    g10["BonusAmount"] = (g10["TotalDelivered"] // 50) * 150\n'
'    results.append(g10)\n'
'\n'
'    result = pd.concat(results, ignore_index=True)\n'
'    return result[\n'
'        result["BonusAmount"] > 0\n'
'    ][["DriverId","AreaId","TotalDelivered","BonusAmount"]]\\\n'
'        .sort_values(["AreaId","BonusAmount"], ascending=[True,False])\\\n'
'        .reset_index(drop=True)')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub(doc, 'ג | תוצאת הרצה על שבוע 15–21 אפריל 2024')

code_block(doc,
'result = calc_weekly_bonus(df, "2024-04-15", "2024-04-21")\n'
'print(result)\n'
'# Output:\n'
'# Empty DataFrame (0 שורות — אף נהג לא הגיע לסף הבונוס)')

data_box(doc, [
    'פלט הפונקציה: DataFrame ריק — 0 שורות.',
    '',
    'הסיבה: ספי הבונוס גבוהים מאוד ביחס לנפח המסירות בדאטאסט:',
    '',
    '  AreaId 1  | סף: 50 מסירות סוף שבוע | מקסימום בפועל בשבוע אחד: 16 מסירות',
    '  AreaId 4  | סף: 300 מסירות שבועיות | מקסימום בפועל בשבוע אחד: 20 מסירות',
    '  AreaId 10 | סף: 50 מסירות סוף שבוע | מקסימום בפועל בשבוע אחד: 12 מסירות',
    '',
    'TOP 3 נהגים קרובים ביותר לבונוס (כל התקופה, לא שבוע בודד):',
    '  AreaId 1:  נהג 188 — 39 מסירות סוף שבוע (78% מהסף)',
    '  AreaId 4:  נהג 844 — 84 מסירות בשבוע    (28% מהסף)',
    '  AreaId 10: נהג 538 — 53 מסירות סוף שבוע (106%!) ← בכל התקופה, לא שבוע',
])

answer(doc, 'הפונקציה עובדת נכון ומיישמת את כל החוקים. עם זאת, בדאטאסט הנוכחי (פברואר–מאי 2024) אף נהג לא עמד בסף הבונוס בשבוע בודד.')
answer_bullet(doc, 'הפונקציה מוכנה לשימוש בכל מאגר נתונים — אם הנפח יגדל, היא תחשב את הבונוסים אוטומטית.')
answer_bullet(doc, 'ממצא חשוב: ספי הבונוס אינם מושגים עם רמות הנפח הנוכחיות — תכנית הבונוסים אינה אפקטיבית במצב הנוכחי.')
answer_bullet(doc, 'המלצה לChecking: ייתכן שהנתונים הם מדגם (28,645 הזמנות ב-18 שבועות = ~236/יום), ובנפח הפרודקשן האמיתי הספים מושגים.')

warning_box(doc,
    'תכנית הבונוסים הנוכחית: אם ספי הבונוס לא מושגים, השליחים לא מקבלים תמריץ. '
    'מומלץ לבצע כיול: AreaId 1/10 — סף של 10–15 מסירות סוף שבוע (ולא 50), '
    'AreaId 4 — סף של 15 מסירות שבועיות (ולא 300). '
    'ניתן לשמור את מבנה הפונקציה ולשנות רק את הפרמטרים.')

insight(doc, 'תובנה: הפונקציה עצמה היא נכס — כלי גנרי לחישוב בונוסים שניתן לעדכן את הספים בלי לשנות קוד. ברגע שנפח HAAT יגדל ×3–5, הפונקציה תתחיל לייצר בונוסים אוטומטית.')

# ══════════════════════════════════════════════════════════════════════════════
#  שמירה
# ══════════════════════════════════════════════════════════════════════════════
out = 'C:/Users/Yosef/PycharmProjects/HAAT_assignment/output/HAAT_DA_Solution.docx'
doc.save(out)
print(f'Saved: {out}')
